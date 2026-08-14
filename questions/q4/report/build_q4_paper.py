"""Build the checked Question 4 Markdown paper as a polished DOCX.

Design preset: narrative_proposal.
Named overrides: A4 Chinese academic page geometry, Chinese typography,
three-line symbol tables, and full-width academic figures.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
REPORT = ROOT / "questions/q4/report"
SOURCE = REPORT / "第四问建模论文.md"
OUTPUT = REPORT / "第四问建模论文.docx"
SUMMARY_PATH = ROOT / "questions/q4/outputs/results/q4_summary.json"
VALIDATION_PATH = ROOT / "questions/q4/outputs/results/q4_validation.json"

CONTENT_WIDTH_DXA = 9072
TABLE_INDENT_DXA = 120
INK = "182B3A"
BLUE = "24557A"
MUTED = "667788"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GRID = "AEB9C4"
GOLD = "A07828"

FORMULA_MAP = {
    1: "X_t^(0) = {商品信息，销售流水，历史批发价格，近期损耗率}                         (1)",
    2: "G = {g_1, g_2, ..., g_13}                                                        (2)",
    3: "g_k = (q_k, n_k, v_k, u_k, s_k, e_k)                                            (3)",
    4: "Y_it = min(D_it, A_it)                                                           (4)",
    5: "L(θ) = ∏_(δ_it=1) f_D(Y_it|X_it;θ) · ∏_(δ_it=0) [1-F_D(Y_it|X_it;θ)]             (5)",
    6: "FR = 1 - [Σ_(i,t)(D_it-Y_it)_+] / [Σ_(i,t)D_it]                                 (6)",
    7: "ln D_ct = α_c + β_c ln P_ct + γ_c^T Z_t + ε_ct                                 (7)",
    8: "W_ibt = I_ibt^beg + Q_ibt - Y_ibt - I_ibt^end - T_ibt                           (8)",
    9: "λ_ibt = W_ibt / (I_ibt^beg + Q_ibt)                                               (9)",
    10: "C_ij = P(i,j) / [P(i)P(j)]                                                      (10)",
    11: "S_ij = P(i被购买|j缺货) - P(i被购买|j有货)                                       (11)",
    12: "D_it = D_c(i)t π_it(A_t)                                                        (12)",
    13: "存在 r∈R_m，使 Σ_(j∈r)y_j = |r|                                                 (13)",
    14: "min Σ_(j∈J)y_j，s.t. 式(13)对所有 m∈K_s 成立                                  (14)",
    15: "对所有 j∈Y_s*，Y_s*\\{j} 均不可行                                              (15)",
    16: "ρ_j = (1/N) Σ_(s=1)^N I(j∈Y_s*)                                                (16)",
    17: "E_j^D = Σ|D_t-D̂_t^(j)|/ΣD_t，E_j^C = Σ|C_t-Ĉ_t^(j)|/ΣC_t                      (17)",
    18: "H_j = |Cov̂_(80%,j) - 0.8|                                                     (18)",
    19: "V_j = (E_0^D-E_j^D, E_0^C-E_j^C, H_0-H_j, P_j-P_0, F_j-F_0, R_j-R_0)            (19)",
    20: "NV_j = N_j(P_j-P_0) - C_j                                                      (20)",
    21: "Y_base* = {库存缺货，批次损耗，供应履约，促销陈列客流，匿名购物篮}               (21)",
}


def set_run_font(run, *, east_asia="宋体", latin="Times New Roman", size=10.5,
                 bold=None, italic=None, color=INK):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, three_line=False):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        enabled = not three_line or edge in {"top", "bottom"}
        node.set(qn("w:val"), "single" if enabled else "nil")
        node.set(qn("w:sz"), "10" if enabled else "0")
        node.set(qn("w:color"), INK if three_line else GRID)
    if three_line:
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_borders = OxmlElement("w:tcBorders")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "7")
            bottom.set(qn("w:color"), INK)
            cell_borders.append(bottom)
            tc_pr.append(cell_borders)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def choose_widths(headers, rows):
    if headers == ["符号", "含义", "单位"]:
        return [1450, 6122, 1500]
    if len(headers) == 4:
        return [1700, 4200, 1650, 1522]
    if len(headers) == 5:
        return [1500, 1100, 1200, 1300, 3972]
    lengths = []
    for idx, header in enumerate(headers):
        values = [len(str(header))] + [len(str(row[idx])) for row in rows]
        lengths.append(max(5, min(max(values), 28)))
    total = sum(lengths)
    widths = [int(CONTENT_WIDTH_DXA * length / total) for length in lengths]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc, headers, rows, *, three_line=False):
    table = doc.add_table(rows=1, cols=len(headers))
    widths = choose_widths(headers, rows)
    set_table_geometry(table, widths)
    set_table_borders(table, three_line=three_line)
    repeat_table_header(table.rows[0])
    font_size = 8.6 if len(headers) >= 5 else 9.0
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        if not three_line:
            set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.05
        set_run_font(p.add_run(normalize_text(str(text))), east_asia="黑体", size=font_size, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in {0, 1, len(headers) - 1} else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            set_run_font(p.add_run(normalize_text(str(text))), size=font_size)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def create_numbering(doc, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "●" if bullet else "%1.")
    level.append(text)
    ppr = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    ppr.append(indent)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 11.5, 8, 4, "1F4D78"),
    ):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
    if "Equation" not in doc.styles:
        equation = doc.styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = doc.styles["Equation"]
    equation.font.name = "Cambria Math"
    equation._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    equation.font.size = Pt(10.0)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(4)
    equation.paragraph_format.space_after = Pt(6)
    equation.paragraph_format.keep_together = True
    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(5)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(paragraph.add_run("第 "), size=9, color=MUTED)
    field = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field._r.extend([begin, instr, end])
    set_run_font(paragraph.add_run(" 页"), size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True
    header = section.header.paragraphs[0]
    set_run_font(header.add_run("2023 C题  |  第四问：数据采集决策与价值验证"), size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])
    props = doc.core_properties
    props.title = "第四问：面向补货与定价改进的数据采集决策模型"
    props.subject = "问题重述、问题分析、模型假设、符号说明、模型建立求解与结果检验"
    props.author = "数学建模项目组"
    props.keywords = "蔬菜；数据采集；删失需求；集合覆盖；Pareto；增量回测"


def add_cover(doc, summary, validation):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("2023 年高教社杯全国大学生数学建模竞赛 C 题"), east_asia="黑体", size=11, bold=True, color=GOLD)
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_before = Pt(70)
    p.add_run("面向补货与定价改进的数据采集决策模型")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    set_run_font(p.add_run("蔬菜类商品的自动定价与补货决策 · 第四问"), size=14, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(32)
    set_run_font(p.add_run("前三问证据审计 + 逻辑能力覆盖 + 最小基数采集组合 + 结构情景灵敏度 + 增量回测-Pareto评价"), size=10.5, color=MUTED)
    rows = [
        ["诊断基础", f"{summary['diagnostic_metric_count']}项前三问证据，{len(summary['source_files'])}个正式源文件"],
        ["候选数据", f"{summary['candidate_data_package_count']}个数据包，{summary['base_capability_count']}项核心能力"],
        ["核心结果", f"唯一{summary['base_portfolio']['package_count']}包最小组合，删一后均不可行"],
        ["复核状态", f"11项单元测试；{validation['passed_count']}/{validation['check_count']}项独立检查通过"],
    ]
    add_table(doc, ["项目", "内容"], rows)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p._p.get_or_add_pPr().append(shd)
    set_run_font(p.add_run("建模边界  "), east_asia="黑体", size=10, bold=True, color=BLUE)
    set_run_font(p.add_run("不为尚未采集的数据虚构精度、先验分布、采集成本或EVSI；事前只判断结构必要性，事后再用同窗口增量回测确认经济价值。"), size=9.5)
    doc.add_page_break()


def normalize_text(text):
    text = text.replace("`", "").replace("**", "").replace("$", "").replace("–", "-").replace("—", "-")
    replacements = {
        r"\mathcal X": "X",
        r"\mathcal G": "G",
        r"\delta": "δ",
        r"\lambda": "λ",
        r"\theta": "θ",
        r"\beta": "β",
        r"\pi": "π",
        r"\widehat": "",
        r"\hat": "",
        r"\bar": "",
        r"\sum": "Σ",
        r"\max": "max",
        r"\mid": "|",
        r"\ge": "≥",
        r"\le": "≤",
        r"\rightarrow": "→",
        r"\in": "∈",
        r"\times": "×",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    text = re.sub(r"\\text\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\^\{([^{}]*)\}", r"^(\1)", text)
    text = re.sub(r"_\{([^{}]*)\}", r"_(\1)", text)
    text = re.sub(r"\\[A-Za-z]+", "", text)
    return text.replace(r"\{", "{").replace(r"\}", "}").replace("\\", "").strip()


def parse_markdown(doc, lines):
    i = 0
    active_num = None
    active_bullet = None
    figure_number = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped.startswith("# ") or stripped.startswith("> "):
            active_num = None
            active_bullet = None
            i += 1
            continue
        if stripped.startswith("## "):
            text = normalize_text(stripped[3:])
            if text in {"参考文献"}:
                doc.add_page_break()
            doc.add_paragraph(text, style="Heading 1")
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(normalize_text(stripped[4:]), style="Heading 2")
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_paragraph(normalize_text(stripped[5:]), style="Heading 3")
            i += 1
            continue
        if stripped == "$$":
            block = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                block.append(lines[i].strip())
                i += 1
            joined = " ".join(block)
            match = re.search(r"\\tag\{(\d+)\}", joined)
            number = int(match.group(1)) if match else None
            p = doc.add_paragraph(style="Equation")
            set_run_font(p.add_run(FORMULA_MAP.get(number, normalize_text(joined))), east_asia="Cambria Math", latin="Cambria Math", size=10.0)
            i += 1
            continue
        image = re.match(r"^!\[(.+?)\]\((.+?)\)$", stripped)
        if image:
            figure_number += 1
            image_path = (REPORT / image.group(2)).resolve()
            if not image_path.is_file():
                raise FileNotFoundError(f"figure not found: {image_path}")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            p.add_run().add_picture(str(image_path), width=Cm(15.8))
            caption = doc.add_paragraph(f"图{figure_number}  {image.group(1)}", style="Caption")
            caption.paragraph_format.keep_with_next = False
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            parsed = [[c.strip() for c in line.strip("|").split("|")] for line in table_lines]
            headers = parsed[0]
            rows = [row for row in parsed[2:] if len(row) == len(headers)]
            add_table(doc, headers, rows, three_line=(headers == ["符号", "含义", "单位"]))
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            if active_num is None or stripped.startswith("1. "):
                active_num = create_numbering(doc, bullet=False)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-0.64)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            apply_numbering(p, active_num)
            p.add_run(normalize_text(numbered.group(1)))
            i += 1
            continue
        bullet = re.match(r"^-\s+(.*)$", stripped)
        if bullet:
            if active_bullet is None:
                active_bullet = create_numbering(doc, bullet=True)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            apply_numbering(p, active_bullet)
            p.add_run(normalize_text(bullet.group(1)))
            i += 1
            continue
        active_num = None
        active_bullet = None
        p = doc.add_paragraph()
        if re.match(r"^\[\d+\]", stripped):
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.65)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(5)
        p.add_run(normalize_text(stripped))
        i += 1


def build():
    for path in (SOURCE, SUMMARY_PATH, VALIDATION_PATH):
        if not path.is_file():
            raise FileNotFoundError(path)
    summary = json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))
    validation = json.loads(VALIDATION_PATH.read_text(encoding="utf-8"))
    if not validation["all_checks_passed"]:
        raise RuntimeError("refusing to build paper from unvalidated Q4 outputs")
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    add_cover(doc, summary, validation)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8").splitlines())
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True
    doc.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    build()
