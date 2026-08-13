"""Build the Question 3 modeling paper DOCX from the checked Markdown source.

Design preset: narrative_proposal.
Named overrides: A4 Chinese academic page geometry, Chinese typography, and a
compact appendix table.  The document is rendered after generation for visual QA.
"""

from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
REPORT = ROOT / "questions" / "q3" / "report"
SOURCE = REPORT / "第三问建模论文.md"
OUTPUT = REPORT / "第三问建模论文.docx"
STRATEGY = ROOT / "questions" / "q3" / "outputs" / "tables" / "q3_daily_strategy.csv"

# A4 academic override: 21.0 x 29.7 cm, 2.5 cm margins.
CONTENT_WIDTH_DXA = 9072
TABLE_INDENT_DXA = 120
INK = "182B3A"
BLUE = "24557A"
MUTED = "667788"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GRID = "AEB9C4"
GOLD = "A07828"
RED = "8A2D2D"


FORMULA_MAP = {
    1: "ŵ_i = w_(i,latest) × ŵ_(c(i),Q2) / w̄_(c(i),recent)                                      (1)",
    2: "w_i^ω = ŵ_i × exp(e_(c(i))^ω)                                                            (2)",
    3: "r̂_i = [Q_i^recent + κ_c r_i^long] / [Q_c^recent + κ_c],   Σ_(i:c(i)=c) r̂_i = 1          (3)",
    4: "r_i^ω = h r_i^(obs,ω) + (1-h) r̂_i,   h = 0.50                                           (4)",
    5: "d_ik^ω = r_i^ω D_(c(i))^ω (p_ik / p_i^ref)^(ε_(c(i)))                                   (5)",
    6: "27 ≤ Σ_(i∈I) x_i ≤ 33                                                                   (6)",
    7: "Σ_(k∈K_i) y_ik = x_i                                                                    (7)",
    8: "2.5 x_i ≤ q_i ≤ M_i x_i                                                                 (8)",
    9: "Σ_(i∈I) q_i ≤ Q^Q2 = 362.262664                                                         (9)",
    10: "0 ≤ s_ik^ω ≤ d_ik^ω y_ik                                                              (10)",
    11: "Σ_(k∈K_i) s_ik^ω ≤ (1-λ_i)q_i                                                         (11)",
    12: "u_c^ω + Σ_(i:c(i)=c) Σ_(k∈K_i) s_ik^ω ≥ D_c^ω,   u_c^ω ≥ 0                            (12)",
    13: "L = [1/(|Ω||C|)] Σ_(ω∈Ω) Σ_(c∈C) u_c^ω / max(D_c^ω, 10⁻⁹)                             (13)",
    14: "L* = min L,   s.t. constraints (6)-(12)                                                (14)",
    15: "Π^ω = Σ_i Σ_k p_ik s_ik^ω - Σ_i w_i^ω q_i                                             (15)",
    16: "L ≤ L* + δ                                                                             (16)",
    17: "ξ_ω ≥ η - Π^ω,   ξ_ω ≥ 0                                                              (17)",
    18: "LTM_α(Π) = η - [1/(α|Ω|)] Σ_(ω∈Ω) ξ_ω                                                 (18)",
    19: "max {(1-γ)E[Π] + γLTM_0.10(Π)},   γ = 0.25                                             (19)",
}


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, latin="Times New Roman", east_asia="宋体", size=10.5,
                 bold=None, color=INK, italic=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    old_grid = table._tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        old_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, *, three_line=False):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        if three_line:
            enabled = edge in {"top", "bottom"}
            node.set(qn("w:val"), "single" if enabled else "nil")
            node.set(qn("w:sz"), "12" if enabled else "0")
            node.set(qn("w:color"), INK)
        else:
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "5")
            node.set(qn("w:color"), GRID)
    if three_line:
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:color"), INK)
            tc_borders.append(bottom)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def create_numbering_definition(doc):
    """Create a real decimal Word numbering sequence that starts at 1."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{(0xA31C0000 + abstract_id) & 0xFFFFFFFF:08X}")
    abstract.append(nsid)
    tmpl = OxmlElement("w:tmpl")
    tmpl.set(qn("w:val"), f"{(0xB42D0000 + abstract_id) & 0xFFFFFFFF:08X}")
    abstract.append(tmpl)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    ppr.append(indent)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def create_bullet_definition(doc):
    """Create a real Word bullet sequence independent of decimal numbering."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{(0xC53E0000 + abstract_id) & 0xFFFFFFFF:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "●")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    ppr.append(indent)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    existing = ppr.find(qn("w:numPr"))
    if existing is not None:
        ppr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)


def column_widths(headers, rows):
    n = len(headers)
    if headers == ["符号", "含义", "单位"]:
        return [1450, 6122, 1500]
    if n == 3:
        return [2100, 4472, 2500]
    if n == 6:
        return [1000, 1430, 1320, 1750, 1750, 1822]
    if n == 8:
        return [1300, 800, 800, 1150, 1150, 1100, 1100, 1672]
    # General deterministic weighting by maximum visible text length.
    lengths = []
    for j, header in enumerate(headers):
        max_len = max([len(str(header))] + [len(str(row[j])) for row in rows])
        lengths.append(max(5, min(max_len, 24)))
    total = sum(lengths)
    widths = [int(CONTENT_WIDTH_DXA * value / total) for value in lengths]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc, headers, rows, *, three_line=False, font_size=8.8, widths=None):
    if len(headers) >= 8:
        font_size = min(font_size, 8.0)
    table = doc.add_table(rows=1, cols=len(headers))
    widths = widths or column_widths(headers, rows)
    set_table_geometry(table, widths)
    set_table_borders(table, three_line=three_line)
    repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        if not three_line:
            set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        set_run_font(p.add_run(str(text)), east_asia="黑体", size=font_size, bold=True)
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_run_font(p.add_run(str(text)), size=font_size, color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def set_paragraph_border(paragraph, color=BLUE, size=12, space=6):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    pbdr.append(left)


def shade_paragraph(paragraph, fill=LIGHT_GRAY):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, east_asia="宋体", size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, east_asia="宋体", size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Times New Roman"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)

    settings = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in settings.items():
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True

    list_style = styles["List Number"]
    list_style.font.name = "Times New Roman"
    list_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    list_style.font.size = Pt(10.5)
    list_style.paragraph_format.left_indent = Cm(0.95)
    list_style.paragraph_format.first_line_indent = Cm(-0.45)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25

    bullet_style = styles["List Bullet"]
    bullet_style.font.name = "Times New Roman"
    bullet_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    bullet_style.font.size = Pt(10.5)
    bullet_style.paragraph_format.left_indent = Cm(1.0)
    bullet_style.paragraph_format.first_line_indent = Cm(-0.45)
    bullet_style.paragraph_format.space_after = Pt(5)
    bullet_style.paragraph_format.line_spacing = 1.25

    if "Equation" not in styles:
        equation = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = styles["Equation"]
    equation.font.name = "Cambria Math"
    equation._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    equation.font.size = Pt(10.5)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(4)
    equation.paragraph_format.space_after = Pt(6)
    equation.paragraph_format.keep_together = True


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(hp.add_run("2023 C题  |  第三问：单品补货与定价联合优化"),
                 east_asia="宋体", size=8.5, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])

    props = doc.core_properties
    props.title = "第三问：销售空间约束下的单品补货与定价联合优化"
    props.subject = "问题重述、问题分析、模型假设、符号说明及模型建立求解"
    props.author = "数学建模项目组"
    props.keywords = "蔬菜；单品补货；定价；MILP；场景优化；经验贝叶斯"


def normalize_text(text):
    text = text.replace("`", "")
    text = text.replace("–", "-").replace("—", "-")
    text = text.replace("×", "×")
    return text.strip()


def add_cover(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("2023 年高教社杯全国大学生数学建模竞赛 C 题"),
                 east_asia="黑体", size=11, bold=True, color=GOLD)

    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_before = Pt(72)
    p.add_run(normalize_text(title))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    set_run_font(p.add_run("蔬菜类商品的自动定价与补货决策"),
                 east_asia="宋体", size=14, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(34)
    set_run_font(p.add_run(normalize_text(subtitle)), east_asia="宋体", size=10.5, color=MUTED)

    rows = [
        ["决策日期", "2023 年 7 月 1 日"],
        ["候选窗口", "2023 年 6 月 24 日-6 月 30 日"],
        ["模型结构", "动态单品份额 + 两阶段词典序场景 MILP"],
        ["复核状态", "17/17 单元测试、77/77 独立输出检查通过"],
    ]
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [1900, 7172])
    set_table_borders(table, three_line=False)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], LIGHT_BLUE)
        for idx, text in enumerate((label, value)):
            pp = cells[idx].paragraphs[0]
            pp.paragraph_format.space_before = Pt(1)
            pp.paragraph_format.space_after = Pt(1)
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(pp.add_run(text), east_asia="黑体" if idx == 0 else "宋体",
                         size=9.5, bold=(idx == 0), color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    shade_paragraph(p, LIGHT_GRAY)
    set_paragraph_border(p, BLUE, 14, 8)
    set_run_font(p.add_run("复核结论  "), east_asia="黑体", size=10, bold=True, color=BLUE)
    set_run_font(p.add_run("模型符合第三问题意，硬约束与会计关系正常；代表场景泛化和六折稳定性达到阈值要求。尾部亏损与单品份额敏感性作为风险边界保留。"),
                 east_asia="宋体", size=9.5, color=INK)
    doc.add_page_break()


def parse_markdown(doc, lines):
    first_title = next(line[2:].strip() for line in lines if line.startswith("# "))
    subtitle = next(line[2:].strip() for line in lines if line.startswith("> "))
    add_cover(doc, first_title, subtitle)

    i = 0
    active_num_id = None
    while i < len(lines):
        raw = lines[i].rstrip("\n")
        stripped = raw.strip()
        if not stripped or stripped.startswith("# ") or stripped.startswith("> "):
            active_num_id = None
            i += 1
            continue
        if stripped.startswith("## "):
            active_num_id = None
            text = normalize_text(stripped[3:])
            if text in {"参考文献", "附录 A：7 月 1 日单品执行清单"}:
                doc.add_page_break()
            doc.add_paragraph(text, style="Heading 1")
            i += 1
            continue
        if stripped.startswith("### "):
            active_num_id = None
            doc.add_paragraph(normalize_text(stripped[4:]), style="Heading 2")
            i += 1
            continue
        if stripped.startswith("#### "):
            active_num_id = None
            doc.add_paragraph(normalize_text(stripped[5:]), style="Heading 3")
            i += 1
            continue
        if stripped == "$$":
            active_num_id = None
            block = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                block.append(lines[i].strip())
                i += 1
            joined = " ".join(block)
            match = re.search(r"\\tag\{(\d+)\}", joined)
            number = int(match.group(1)) if match else None
            text = FORMULA_MAP.get(number, normalize_text(joined))
            p = doc.add_paragraph(style="Equation")
            p.add_run(text)
            i += 1
            continue
        if stripped.startswith("|"):
            active_num_id = None
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            parsed = [[cell.strip() for cell in line.strip("|").split("|")] for line in table_lines]
            headers = parsed[0]
            rows = [row for row in parsed[2:] if len(row) == len(headers)]
            add_table(doc, headers, rows, three_line=(headers == ["符号", "含义", "单位"]))
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            if active_num_id is None or stripped.startswith("1. "):
                active_num_id = create_numbering_definition(doc)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-0.64)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            apply_numbering(p, active_num_id)
            p.add_run(normalize_text(numbered.group(1)))
            i += 1
            continue
        bullet = re.match(r"^-\s+(.*)$", stripped)
        if bullet:
            active_num_id = None
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(normalize_text(bullet.group(1)))
            i += 1
            continue
        active_num_id = None
        p = doc.add_paragraph()
        if stripped.startswith("[1]") or stripped.startswith("[2]") or stripped.startswith("[3]"):
            p.paragraph_format.first_line_indent = Cm(-0.65)
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(5)
        p.add_run(normalize_text(stripped))
        i += 1


def add_appendix_strategy(doc):
    with STRATEGY.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = [row for row in csv.DictReader(handle) if int(float(row["selected"])) == 1]
    category_order = ["水生根茎类", "花叶类", "花菜类", "茄类", "辣椒类", "食用菌"]
    counter = 1
    for category in category_order:
        part = [row for row in rows if row["category_name"] == category]
        p = doc.add_paragraph(f"附录 A.{category_order.index(category)+1} {category}", style="Heading 2")
        p.paragraph_format.keep_with_next = True
        table_rows = []
        for row in part:
            table_rows.append([
                str(counter),
                row["sku_name"],
                f'{float(row["price_yuan_per_kg"]):.1f}',
                f'{float(row["order_qty_kg"]):.3f}',
                f'{float(row["expected_sales_kg"]):.3f}',
                f'{100*float(row["stockout_probability"]):.1f}%',
                f'{float(row["expected_profit_yuan"]):.2f}',
            ])
            counter += 1
        add_table(
            doc,
            ["序号", "单品名称", "售价\n(元/kg)", "订货量\n(kg)", "期望销量\n(kg)", "缺货概率", "期望利润\n(元)"],
            table_rows,
            font_size=8.2,
            widths=[1000, 1800, 1000, 1250, 1250, 1100, 1672],
        )


def build_up_equations_with_word(path):
    """Convert linear equation paragraphs to native Word equations when available."""
    try:
        import win32com.client
    except ImportError:
        print("pywin32 unavailable: keeping readable linear equations")
        return
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    document = None
    converted = 0
    try:
        document = word.Documents.Open(str(path.resolve()))
        for paragraph in document.Paragraphs:
            style_name = str(paragraph.Style.NameLocal)
            if style_name not in {"Equation", "公式"}:
                continue
            rng = paragraph.Range.Duplicate
            rng.End = max(rng.Start, rng.End - 1)
            if not rng.Text.strip():
                continue
            document.OMaths.Add(rng)
            document.OMaths(document.OMaths.Count).BuildUp()
            converted += 1
        document.Save()
    except Exception as error:
        print(f"Word equation build-up skipped: {error}")
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()
    print(f"built {converted} native Word equations")


def restore_numbering_sequences(path):
    """Reapply three independent real-numbered lists after Word equation conversion."""
    doc = Document(path)
    groups = [[
        "读取附件处理数据", "估计单品成本", "将份额场景", "从 600 场景中分层",
        "把每个 K 的决策", "生成单品清单", "运行 17 项单元测试",
    ]]
    for prefixes in groups:
        num_id = create_numbering_definition(doc)
        found = []
        for prefix in prefixes:
            match = next((p for p in doc.paragraphs if p.text.startswith(prefix)), None)
            if match is None:
                raise RuntimeError(f"numbered paragraph not found: {prefix}")
            apply_numbering(match, num_id)
            found.append(match)
        if len(found) != len(prefixes):
            raise RuntimeError("numbering sequence restoration was incomplete")
    bullet_groups = [
        [
            "题意与数据口径", "目标优先级符合", "能处理需求", "复杂度与可执行性",
            "创新建立在现有数据",
        ],
        ["单品需求份额", "模型未显式", "极端场景仍存在"],
    ]
    for prefixes in bullet_groups:
        num_id = create_bullet_definition(doc)
        for prefix in prefixes:
            match = next((p for p in doc.paragraphs if p.text.startswith(prefix)), None)
            if match is None:
                raise RuntimeError(f"bullet paragraph not found: {prefix}")
            apply_numbering(match, num_id)
    doc.save(path)
    print("restored Word numbering and bullet sequences")


def build():
    if not SOURCE.exists() or not STRATEGY.exists():
        raise FileNotFoundError("paper source or Q3 strategy output is missing")
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    parse_markdown(doc, lines)
    add_appendix_strategy(doc)

    # Prevent accidental widows/orphans for body text.
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True
    doc.save(OUTPUT)
    build_up_equations_with_word(OUTPUT)
    restore_numbering_sequences(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    build()
